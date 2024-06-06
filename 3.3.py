from docx import Document

doc = Document('5.docx')
out = []
stroka = ""
for i in doc.paragraphs:
    temp = i.text
    count = 0
    for j in temp:
        if j == ",":
            count += 1
            break
    if count > 0:
        out.append(temp + " | True 1 ---> В данной строке есть запятые" + "\n")
        stroka += "1"
    else:
        out.append(temp + " | False 0 ---> В данной строке нет запятых" + "\n")
        stroka += "0"
    count = 0
print(*out, sep="\n")
doc = Document()
for i in out:
    doc.add_paragraph(i)
doc.save('15_itog.docx')
print(stroka)
doc = Document()
doc.add_paragraph(stroka)
doc.save('15_code.docx')
