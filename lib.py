from docx import Document
import codecs
# from docx_utils.flatten import opc_to_flat_opc
from test import findersomeone, atomikos
import MTK2

doc = Document('variant15.docx')
# opc_to_flat_opc('variant15.docx', 'variant15.xml')

m1 = doc.paragraphs[0].runs[0].font.size
m2 = doc.paragraphs[0].runs[0].font.color.rgb
m3 = doc.paragraphs[0].runs[0].font.highlight_color

for suspect in range(4):
    temp = []
    for p in doc.paragraphs:
        for run in p.runs:
            print(len(run.text), run.text, run.font.size, run.font.color.rgb, run.font.highlight_color)
            if suspect == 0:
                if m1 == run.font.size:
                    temp.append("1" * len(run.text))
                else:
                    temp.append("0" * len(run.text))
            if suspect == 1:
                if m2 == run.font.color.rgb:
                    temp.append("1" * len(run.text))
                else:
                    temp.append("0" * len(run.text))
            if suspect == 2:
                if m3 == run.font.highlight_color:
                    temp.append("1" * len(run.text))
                else:
                    temp.append("0" * len(run.text))

    out = ""
    for i in temp:
        out = out + i
    for suspect in range(3):
        output = ""
        for i in range(len(out) // 8):
            # print(out[0+i*8:8+i*8])
            bytes_hex = hex(int(out[0 + i * 8:8 + i * 8], 2))
            # print(bytes_hex[2:4])
            if (bytes_hex != 0) & (len(bytes_hex) > 3):
                if suspect == 0:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='koi8-r')
                if suspect == 1:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp866')
                if suspect == 2:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp1251')
        print(output)
    output = MTK2.MTK2_decode(out)
    print(output)

out = atomikos("result.xml")
# out = somik("variant1.xml")
# print(out)
for suspect in range(3):
    output = ""
    for i in range(len(out) // 8):
        # print(out[0+i*8:8+i*8])
        bytes_hex = hex(int(out[0 + i * 8:8 + i * 8], 2))
        # print(bytes_hex[2:4])
        if (bytes_hex != 0) & (len(bytes_hex) > 3):
            if suspect == 0:
                output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='koi8-r')
            if suspect == 1:
                output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp866')
            if suspect == 2:
                output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp1251')
    print(output)
output = MTK2.MTK2_decode(out)
print(output)
# print(huh.decode('koi8-r'))
# print(huh.decode('cp866'))
# print(huh.decode('cp1251'))
doc.save('test2.docx')

# lol = "sos"
# print(type(lol))
# koi = lol.encode(encoding='koi8-r',errors='ignorea')
# print(koi)
# print(koi.decode('koi8-r'))
