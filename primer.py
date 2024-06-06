import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

doc = docx.Document('5.docx')

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


def run_set_spacing(run, value: int):
    """Set the font spacing for `run` to `value` in twips.

    A twip is a "twentieth of an imperial point", so 1/1440 in.
    """

    def get_or_add_spacing(rPr):
        # --- check if `w:spacing` child already exists ---
        spacings = rPr.xpath("./w:spacing")
        # --- return that if so ---
        if spacings:
            return spacings[0]
        # --- otherwise create one ---
        spacing = OxmlElement("w:spacing")
        rPr.insert_element_before(
            spacing,
            *(
                "w:w",
                "w:kern",
                "w:position",
                "w:sz",
                "w:szCs",
                "w:highlight",
                "w:u",
                "w:effect",
                "w:bdr",
                "w:shd",
                "w:fitText",
                "w:vertAlign",
                "w:rtl",
                "w:cs",
                "w:em",
                "w:lang",
                "w:eastAsianLayout",
                "w:specVanish",
                "w:oMath",
            ),
        )
        return spacing

    rPr = run._r.get_or_add_rPr()
    spacing = get_or_add_spacing(rPr)
    spacing.set(qn('w:val'), str(value))


text = "Трудовое добро ни в воде не тонет, ни на огне не горит."


def encode_to_binary(text):
    encodedd = [bin(ord(char))[2:].zfill(16) for char in text]
    # преобразуем список двоичных чисел в одну строку
    binary_string = ''.join(encodedd)
    return binary_string


ascii = encode_to_binary(text)
print(ascii)
LenParagraphs = []
OpenText = ""
for paragraph in doc.paragraphs:
    stroka = ""
    for run in paragraph.runs:
        for char in run.text:
            stroka += char
    OpenText += stroka
    LenParagraphs.append(len(stroka))
print(LenParagraphs)
print(OpenText)

doc.paragraphs.clear()

id_char = 0
for i in range(len(doc.paragraphs)):  # прогоняем все строки
    doc.paragraphs[i].clear()
    for id_rans in range(LenParagraphs[i]):  # прогоняем каждую строку
        run = doc.paragraphs[i].add_run(OpenText[id_char])  # добавляем заново текст
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.highlight_color = WD_COLOR_INDEX.WHITE
        if id_char < len(ascii):
            if ascii[id_char] == '1':
                run_set_spacing(run, 10)

            else:
                pass

        else:
            pass

        id_char += 1

doc.save('result.docx')
